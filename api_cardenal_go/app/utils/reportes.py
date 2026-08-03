from io import BytesIO
from fastapi.responses import StreamingResponse
from geoalchemy2.elements import WKBElement
from geoalchemy2.shape import to_shape

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


def obtenerColumnas(item) -> list:
    if isinstance(item, dict):
        return list(item.keys())
    elif hasattr(item, '__table__'):
        return [c.name for c in item.__table__.columns]
    return []


def obtenerValorLimpio(obj, campo: str) -> str:
    valor = obj.get(campo) if isinstance(obj, dict) else getattr(obj, campo, None)
    if valor is None:
        return "N/A"
    if isinstance(valor, WKBElement):
        try:
            shape_point = to_shape(valor)
            return f"POINT({shape_point.x:.5f} {shape_point.y:.5f})"
        except Exception:
            return "Geometría Inválida"
    return str(valor)


def generarReporteWord(lista_datos: list, titulo: str) -> StreamingResponse:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = heading.add_run(titulo.upper())
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 41, 59)

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("Reporte generado automáticamente por el sistema Cardenal Go")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(9)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    p_sub.paragraph_format.space_after = Pt(14)

    if lista_datos:
        columnas = obtenerColumnas(lista_datos[0])
        tabla = doc.add_table(rows=1, cols=len(columnas))
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.autofit = True

        hdr_cells = tabla.rows[0].cells
        for i, col in enumerate(columnas):
            hdr_cells[i].text = col.replace("_", " ").title()
            shading_elm = parse_xml(r'<w:shd {} w:fill="1E293B"/>'.format(nsdecls('w')))
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, item in enumerate(lista_datos):
            row_cells = tabla.add_row().cells
            bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            
            for i, col in enumerate(columnas):
                row_cells[i].text = obtenerValorLimpio(item, col)
                if bg_color != "FFFFFF":
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
                    row_cells[i]._tc.get_or_add_tcPr().append(shd)

                p = row_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(51, 65, 85)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        headers={"Content-Disposition": f"attachment; filename={titulo}.docx"}
    )


def generarReportePDF(lista_datos: list, titulo: str) -> StreamingResponse:
    buffer = BytesIO()
    columnas = obtenerColumnas(lista_datos[0]) if lista_datos else []
    
    pagesize = landscape(letter) if len(columnas) > 5 else letter

    doc = SimpleDocTemplate(
        buffer,
        pagesize=pagesize,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=4
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=8,
        leading=10,
        textColor=colors.HexColor('#64748B'),
        spaceAfter=14
    )

    header_cell_style = ParagraphStyle(
        'HeaderCell',
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=11,
        textColor=colors.white,
        alignment=1
    )

    body_cell_style = ParagraphStyle(
        'BodyCell',
        fontName='Helvetica',
        fontSize=8,
        leading=10,
        textColor=colors.HexColor('#334155'),
        alignment=0
    )

    elements = [
        Paragraph(titulo.upper(), title_style),
        Paragraph("Sistema Cardenal Go - Reporte Automatizado", subtitle_style)
    ]

    if lista_datos:
        data = []
        headers = [Paragraph(col.replace("_", " ").title(), header_cell_style) for col in columnas]
        data.append(headers)

        for item in lista_datos:
            row = []
            for col in columnas:
                val = obtenerValorLimpio(item, col)
                row.append(Paragraph(val, body_cell_style))
            data.append(row)

        t_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E293B')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ])

        for i in range(1, len(data)):
            if i % 2 == 0:
                t_style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor('#F8FAFC'))

        tabla = Table(data, repeatRows=1)
        tabla.setStyle(t_style)
        elements.append(tabla)

    doc.build(elements)
    buffer.seek(0)
    return StreamingResponse(
        buffer, 
        media_type="application/pdf", 
        headers={"Content-Disposition": f"attachment; filename={titulo}.pdf"}
    )


def generarReporteExcel(lista_datos: list, titulo: str) -> StreamingResponse:
    wb = Workbook()
    ws = wb.active
    ws.title = titulo[:30]

    ws.views.sheetView[0].showGridLines = True

    if lista_datos:
        columnas = obtenerColumnas(lista_datos[0])

        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(columnas))
        title_cell = ws.cell(row=1, column=1, value=titulo.upper())
        title_cell.font = Font(name='Calibri', size=14, bold=True, color='1E293B')
        title_cell.alignment = Alignment(horizontal='left', vertical='center')
        ws.row_dimensions[1].height = 25

        header_row = 3
        headers = [col.replace("_", " ").title() for col in columnas]
        
        header_fill = PatternFill(start_color='1E293B', end_color='1E293B', fill_type='solid')
        header_font = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
        header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        thin_border = Border(
            left=Side(style='thin', color='CBD5E1'),
            right=Side(style='thin', color='CBD5E1'),
            top=Side(style='thin', color='CBD5E1'),
            bottom=Side(style='thin', color='CBD5E1')
        )

        for col_num, header_title in enumerate(headers, 1):
            cell = ws.cell(row=header_row, column=col_num, value=header_title)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment
            cell.border = thin_border

        ws.row_dimensions[header_row].height = 24

        alt_fill = PatternFill(start_color='F8FAFC', end_color='F8FAFC', fill_type='solid')
        data_font = Font(name='Calibri', size=10, color='334155')
        data_alignment = Alignment(horizontal='left', vertical='center')

        for r_idx, item in enumerate(lista_datos, start=header_row + 1):
            row_data = [obtenerValorLimpio(item, col) for col in columnas]
            for c_idx, val in enumerate(row_data, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=val)
                cell.font = data_font
                cell.alignment = data_alignment
                cell.border = thin_border
                if (r_idx - header_row) % 2 == 0:
                    cell.fill = alt_fill
            ws.row_dimensions[r_idx].height = 20

        for col in ws.columns:
            max_len = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col[2:]:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = max(max_len + 4, 12)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer, 
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
        headers = {"Content-Disposition": f"attachment; filename={titulo}.xlsx"}
    )